from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\laragon\www\pms_systemv2")
ASSET = ROOT / "manual_assets"
OUT = ROOT / "PMAMS_Non-Technical_User_Manual_v2.docx"


def set_font(run, name="Calibri", size=10.5, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7C5D6", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, widths, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        if header and row_idx == 0:
            set_repeat_header(row)
        for col_idx, cell in enumerate(row.cells):
            if col_idx < len(widths):
                set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        set_font(run, size=9, bold=True, color=(31, 78, 121))
            else:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_font(run, size=9)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(46, 116, 181))
    elif level == 2:
        set_font(run, size=13, bold=True, color=(46, 116, 181))
    else:
        set_font(run, size=11.5, bold=True, color=(31, 78, 121))
    return p


def add_para(doc, text="", bold_prefix=None, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=10.5, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, size=10.5, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_font(r, size=10.5, italic=italic, color=color)
    return p


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(item)
        set_font(r, size=10.5)


def add_screenshot(doc, filename, caption, width=6.2):
    path = ASSET / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption)
    set_font(r, size=8.5, italic=True, color=(89, 101, 114))


def add_callout(doc, title, text, fill="EEF5FB", border="6EA8DC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.25)
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    set_cell_shading(cell, fill)
    set_table_borders(table, border, "10")
    p = cell.paragraphs[0]
    r = p.add_run(title + " ")
    set_font(r, size=10.5, bold=True, color=(31, 78, 121))
    r = p.add_run(text)
    set_font(r, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8, color=(89, 101, 114))
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Calibri"
        styles[style_name]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        styles[style_name]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    # Header and footer
    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("PMAMS  |  Non-technical User Manual")
    set_font(r, size=8.5, bold=True, color=(46, 116, 181))
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = fp.add_run("PMAMS • Preventive Maintenance and Asset Monitoring System")
    set_font(r, size=8, color=(89, 101, 114))
    fp.add_run("    ")
    set_page_number(fp)

    doc.core_properties.title = "PMAMS Non-technical User Manual"
    doc.core_properties.subject = "User guide for the Preventive Maintenance and Asset Monitoring System"
    doc.core_properties.author = "PMAMS"
    doc.core_properties.keywords = "PMAMS, equipment, preventive maintenance, user manual"
    return doc


def build():
    doc = setup_document()

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run()
    r.add_picture(str(ASSET / "ictu-logo.png"), width=Inches(1.4))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PMAMS")
    set_font(r, size=30, bold=True, color=(31, 78, 121))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Preventive Maintenance and Asset Monitoring System")
    set_font(r, size=17, bold=True, color=(46, 116, 181))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Non-technical user manual")
    set_font(r, size=13, italic=True, color=(89, 101, 114))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("For administrators, unit heads/deans, custodians, and staff")
    set_font(r, size=11.5, color=(64, 74, 84))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version 3.5.7  |  August 2026")
    set_font(r, size=10, color=(89, 101, 114))
    add_callout(doc, "Purpose.", "This guide explains everyday tasks in plain language. It intentionally avoids programming and server administration details.")
    doc.add_page_break()

    add_heading(doc, "Contents", 1)
    contents = [
        "1. Before you begin",
        "2. Sign in and understand your role",
        "3. Navigate PMAMS",
        "4. Manage locations, offices, and staff",
        "5. Register and maintain equipment",
        "6. Issue, return, and reissue equipment",
        "7. Complete a maintenance checklist",
        "8. Set up and monitor preventive maintenance plans",
        "9. Use reports and exports",
        "10. Maintenance Attention recommendations",
        "11. Maintenance Photo Gallery",
        "12. User accounts, privacy, and activity logs",
        "13. Backup, restore, and deployment checks",
        "14. Troubleshooting and quick reference",
    ]
    add_bullets(doc, contents)
    add_para(doc, "Tip: Use the headings and the sidebar names in the application together. Menu entries can differ by role because Super Admin controls role-based access.", italic=True, color=(89, 101, 114))

    add_heading(doc, "1. Before you begin", 1)
    add_para(doc, "PMAMS is a browser-based system for recording ICT equipment, locations, staff assignments, preventive maintenance checklists, reports, and maintenance recommendations.")
    add_table(doc, ["You need", "Why it matters"], [
        ["A supported browser", "Use an up-to-date Chrome, Edge, Firefox, or the mobile browser provided by your organization."],
        ["Your assigned account", "Your role determines which menus and actions are available."],
        ["A reachable PMAMS address", "Use the official HTTPS address when deployed. The local XAMPP/Laragon address is for internal testing."],
        ["Camera permission (optional)", "Needed only when taking equipment or maintenance photos and scanning QR codes."],
    ], [2.0, 4.25])
    add_callout(doc, "Important.", "Never share your password. Use Remember me only on a private, trusted device. Sign out on shared computers.", fill="FFF6E5", border="E5A52B")

    add_heading(doc, "2. Sign in and understand your role", 1)
    add_para(doc, "Open the PMAMS login page, enter your email and password, and select Sign in. If Remember me is enabled, the browser may keep your session for the configured period; this does not replace normal security controls.")
    add_bullets(doc, [
        "Super Admin: full system control, including users, role permissions, recycle bin, backup/restore, PM Plan setup, and all reports.",
        "Admin: manages equipment, locations, staff, reports, and assigned maintenance work. Access to users and destructive cleanup depends on role settings.",
        "Unit Head/Dean: reviews or signs assigned work and can see the locations or college data assigned to the account.",
        "Custodian: manages permitted locations, offices, staff, and equipment activities. Marking maintenance is controlled separately by role permissions.",
    ])
    add_table(doc, ["If you need to…", "Start here"], [
        ["Add or locate equipment", "Equipment"],
        ["Assign an office or staff member", "Locations → office → staff"],
        ["Schedule preventive maintenance", "PM Plan (Super Admin)"],
        ["Mark equipment checked", "Equipment → Mark checked (when enabled)"],
        ["Print or export records", "Reports"],
    ], [2.5, 3.75])
    add_screenshot(doc, "mobile-menu.png", "Figure 1. The mobile sidebar. The visible menu depends on the signed-in role.", 3.0)

    add_heading(doc, "3. Navigate PMAMS", 1)
    add_para(doc, "The left sidebar is the main navigation. On a phone, use the hamburger button to open it. The header shows the current page and signed-in account. The light/dark theme control changes appearance only; it does not change data.")
    add_bullets(doc, [
        "Dashboard: quick counts, charts, recent activity, PM status, and shortcuts.",
        "Locations: registered locations, offices, staff, and equipment assigned to each office.",
        "Equipment: inventory, filters, add/edit, photo capture, linking, issuance, history, QR, and deletion actions.",
        "PM Plan: schedules and completion monitoring for preventive maintenance.",
        "Reports: filtered views and printable PDF/Excel outputs.",
        "PM Gallery: maintenance photographs with date filtering and slideshow/grid views.",
        "Activity Logs: audit trail of important system actions.",
    ])
    add_callout(doc, "Page loading behavior.", "Some large pages intentionally stay empty until you press Apply/Generate or Reset. This prevents the browser from loading every record at once. Keep your typed search text and press the action button when available.")

    add_heading(doc, "4. Manage locations, offices, and staff", 1)
    add_heading(doc, "4.1 Add a location", 2)
    add_bullets(doc, [
        "Open Locations and choose Add Location.",
        "Enter the location or college name and save.",
        "Use the location filter first when a long list is displayed.",
        "Select the location name to manage its offices.",
    ], numbered=True)
    add_heading(doc, "4.2 Add an office and designate a head/dean", 2)
    add_bullets(doc, [
        "Open a location, then choose Add Office.",
        "Enter the office name. Office is a child of Location, so select the location before filtering offices.",
        "Open the office staff list and add or edit a staff record.",
        "Use the role/position or designation field for the person who should appear as Head of Unit or Dean in reports.",
        "Only one designated head/dean should be used for an office at a time. Update the designation when responsibility changes.",
    ], numbered=True)
    add_screenshot(doc, "offices.png", "Figure 2. Office list under a registered location.", 6.0)
    add_screenshot(doc, "staff.png", "Figure 3. Staff records for an office.", 6.0)
    add_callout(doc, "Good practice.", "Create the location and office master data before importing equipment. This makes office filters, issuance, PM schedules, and report signatories reliable.")

    add_heading(doc, "5. Register and maintain equipment", 1)
    add_heading(doc, "5.1 Add equipment", 2)
    add_para(doc, "Open Equipment and choose Add Equipment. The exact fields depend on the equipment type. Desktop and Laptop records may include computer name, memory, storage, operating system, and office software. Printers, monitors, UPS/AVR, scanners, and network devices show the fields relevant to them.")
    add_bullets(doc, [
        "Use a unique Property Number when one is available.",
        "For a peripheral that belongs to a system unit, use Part of Property Number to link the child device to the parent property number.",
        "If the property number is blank, PMAMS can generate a temporary series such as DESKTOP-TempID-YYYYMMDD-####. Keep the generated value; do not manually reuse it.",
        "Use the Memory and Storage selections consistently. Storage capacity appears after the storage type is selected.",
        "Select Condition and Status carefully. These values affect dashboard counts, maintenance attention, and reports.",
        "Use Take Photo on a device with camera permission. The image is stored with the equipment record and can be cleared later.",
    ])
    add_heading(doc, "5.2 Search and filter inventory", 2)
    add_para(doc, "The Equipment page uses a filter row for equipment type, location, status, condition, and a text search. Search can include property number, serial number, computer name, brand, or model. In a large deployment, press Apply or Reset to load the matching data.")
    add_screenshot(doc, "equipment-filters.png", "Figure 4. Equipment search and filter controls.", 6.2)
    add_heading(doc, "5.3 Edit, link, history, and delete", 2)
    add_bullets(doc, [
        "View opens the equipment detail page.",
        "Link connects a peripheral to a parent system unit. The parent property number remains the reference for the group.",
        "History shows issuance, relocation, reissue, checklist, and corrective-action events.",
        "Edit Specs updates the equipment record. The equipment type is locked after creation to protect historical reporting.",
        "Delete moves the record to the recycle bin when soft-delete protection is enabled. Permanent deletion cannot be undone and is restricted to authorized users.",
    ])
    add_callout(doc, "Bulk deletion.", "Use the page checkbox or the all-filtered checkbox only after checking the matching count. Destructive actions should be confirmed and, where configured, are recorded in Activity Logs.", fill="FFF6E5", border="E5A52B")

    add_heading(doc, "6. Issue, return, and reissue equipment", 1)
    add_para(doc, "Issuance connects an equipment record to an active staff member. Location and office context come from the staff member's registered office.")
    add_bullets(doc, [
        "Open the staff member or equipment detail page and choose Issue/Reissue when available.",
        "Type at least two characters in the staff search. Select the registered person from the suggestions; do not create a free-text name.",
        "The equipment remains uniquely identified by its property number. Reissue changes the assigned end user and records the movement.",
        "Use Return when the item is no longer assigned. Confirm the activity before submitting.",
        "If the staff member is missing, use Add Staff only when your role permits it, then return to the issuance form.",
    ])
    add_screenshot(doc, "staff-equipment.png", "Figure 5. Equipment currently assigned to a staff member.", 6.2)
    add_callout(doc, "Audit trail.", "Use a short, clear remark for unusual movements—for example, “Transferred to new office” or “Returned for repair.” The remark becomes useful history later.")

    add_heading(doc, "7. Complete a maintenance checklist", 1)
    add_para(doc, "A checklist is used to record the preventive maintenance result for eligible equipment. The page may show linked peripheral rows such as monitor, printer, UPS, or AVR.")
    add_heading(doc, "7.1 Before saving", 2)
    add_bullets(doc, [
        "Verify the office/unit, college/location, maintenance date, and system unit property number.",
        "Choose one outcome for every checklist item: OK or Not OK. UPS/AVR and Printer may also offer Not Available when configured.",
        "If Not OK is selected, choose the relevant condition/status such as Repair, Condemned, or Not in Use when shown.",
        "If Not Available is selected, the peripheral can be unlinked and the property number may not be required for that row.",
        "Review the generated remarks and corrective action. Add a reason when the duplicate-check window asks for verification.",
        "Add a maintenance photo with Take photo or Upload photo when the module is enabled. Allow camera permission on mobile.",
        "Save only after all required rows and fields are complete. A completion message confirms that the record was saved.",
    ])
    add_callout(doc, "Duplicate check.", "If a device was checked recently within the configured window, PMAMS asks whether you want to verify the checklist. This does not erase the earlier historical record.", fill="FFF6E5", border="E5A52B")
    add_heading(doc, "7.2 Linked peripherals", 2)
    add_bullets(doc, [
        "Use the link icon beside a Not linked peripheral to search available equipment.",
        "Select the correct peripheral and save the link. The checklist page should retain current selections while the link is saved.",
        "Use Unlink only when the peripheral no longer belongs to the parent system unit.",
        "A checklist photo saved from the checklist is also available in PM Gallery when the record is stored successfully.",
    ])
    add_screenshot(doc, "maintenance-table.png", "Figure 6. A maintenance/report table showing recorded checklist results.", 6.2)

    add_heading(doc, "8. Set up and monitor preventive maintenance plans", 1)
    add_para(doc, "PM Plan schedules maintenance by location and office. The Super Admin publishes the original schedule; assigned administrators work within their allowed locations.")
    add_heading(doc, "8.1 Create and publish a plan", 2)
    add_bullets(doc, [
        "Open PM Plan and choose Add or Schedule PM Plan.",
        "Select the location first, then select one or more offices. Office choices depend on the selected location.",
        "Choose a month or exact maintenance date, then assign one or more administrators where permitted.",
        "Review duplicate warnings for the same location, office, and date before publishing.",
        "Publish the plan. A success message confirms the schedule is available to assigned users.",
        "Use pagination to review all schedules; do not assume the first page is the complete list.",
    ])
    add_screenshot(doc, "pm-plan-card.png", "Figure 7. PM Plan schedule card with progress and schedule actions.", 6.2)
    add_heading(doc, "8.2 Temporary override", 2)
    add_para(doc, "An assigned administrator can override a schedule when an office proposes a reschedule. The original schedule stays in the record; the override appears beside it with the reason.")
    add_bullets(doc, [
        "Open Override schedule on the plan card.",
        "Enter the exact override date and a concise reason/remarks.",
        "Save the override. Use Reset/remove override only when authorized and certain the temporary schedule should be cleared.",
        "The monitoring report displays the original schedule, override date, and reason when applicable.",
    ])
    add_screenshot(doc, "pm-override.png", "Figure 8. Override schedule form with reason and reset/remove action.", 4.5)
    add_heading(doc, "8.3 Record completion details", 2)
    add_bullets(doc, [
        "When all eligible equipment for an office is checked, record completion details if the form appears.",
        "The person conducting the maintenance is taken from the checklist records. Add the signer name and use the on-screen signature pad when requested.",
        "The monitoring report uses the completion date, person in charge, signature, and remarks.",
    ])

    add_heading(doc, "9. Use reports and exports", 1)
    add_para(doc, "Reports are designed to be filtered before loading large result sets. Select filters and press Generate, Filter, Apply, or Reset as displayed on the page.")
    add_table(doc, ["Report", "Use it for", "Helpful filters"], [
        ["All Assets", "Complete inventory export, including equipment that is not issued.", "Location, office, type, status, condition, text search"],
        ["Issuance", "Current and historical assignments, returns, and transfers.", "Staff, equipment type, location, office, semester/date"],
        ["Checked Equipment", "Saved checklist history and printable checklist records.", "Date range, location, property/remark search"],
        ["PM Schedule Monitoring", "Compare scheduled maintenance with actual dates and sign-off.", "Location, office, schedule date"],
        ["PM Quality Objective", "Compare baseline and actual maintenance counts and status.", "Year, semester, location, office"],
        ["Maintenance Attention", "Find equipment needing upgrade, license, storage, RAM, or age attention.", "Year, semester, location, office, type, priority, condition, status"],
    ], [1.6, 2.7, 1.95])
    add_callout(doc, "PDF and Excel.", "Use Preview to check the filtered result. Use Print/Open PDF for the formatted PDF and Export for Excel. The report should use the same filter set in both formats.")
    add_screenshot(doc, "maintenance-attention-report.png", "Figure 9. Maintenance Attention report with priority, score, recommendation, and PM schedule.", 6.2)
    add_heading(doc, "9.1 Location and office filters", 2)
    add_bullets(doc, [
        "Select a location first. The Office list then refreshes to show only offices belonging to that location.",
        "Choose All offices to include every office under the selected location.",
        "If a record is missing, clear the text search, choose the correct location, and press Apply/Generate again.",
    ])
    add_screenshot(doc, "offices.png", "Figure 10. Offices are managed under their parent location.", 5.5)

    add_heading(doc, "10. Maintenance Attention recommendations", 1)
    add_para(doc, "Maintenance Attention is a planning aid. It does not replace a technician's inspection. It applies the selected rules or the configured local model mode to eligible equipment and shows a priority, score, and recommended action.")
    add_heading(doc, "10.1 Current rule examples", 2)
    add_bullets(doc, [
        "Windows 10/11 with RAM at or below 8 GB → consider upgrading RAM to at least 16 GB.",
        "Windows 8 or older with RAM at or below 4 GB → consider upgrading RAM to at least 8 GB.",
        "HDD storage on Desktop → consider replacing or upgrading to SSD flash storage.",
        "Cracked or non-genuine operating system or Microsoft Office → procure a genuine license.",
        "Equipment procured at least six years ago → review for replacement or major upgrade.",
        "Condition and status indicators contribute to priority. Condemned items are not counted as maintained equipment.",
    ])
    add_screenshot(doc, "maintenance-attention-types.png", "Figure 11. Equipment-type multi-selection in Maintenance Attention.", 3.2)
    add_para(doc, "The page can show the last local model training date/time when local AI mode is enabled. If Python is unavailable, use the Laravel coded rules mode; the system remains usable without a model.", italic=True, color=(89, 101, 114))

    add_heading(doc, "11. Maintenance Photo Gallery", 1)
    add_para(doc, "PM Gallery stores photographs taken during preventive maintenance. A photo may be linked to equipment or saved as a general maintenance photo when the form allows it.")
    add_bullets(doc, [
        "Use Take photo on a phone to open the rear camera. Allow camera permission the first time.",
        "Use Upload photo to select an existing image. Photos are limited to the configured file size.",
        "Filter by capture date, property number, serial number, type, or caption.",
        "Use Grid for quick review and Auto slide for a hands-free presentation.",
        "Select photos for bulk download (ZIP) or deletion. Users may delete only photos they uploaded unless a higher role grants broader access.",
    ])
    add_callout(doc, "If the camera does not open.", "Use HTTPS, confirm browser camera permission, close another app using the camera, and reload the page. On desktop, the browser may offer a file picker instead of a camera if no camera is available.")

    add_heading(doc, "12. User accounts, privacy, and activity logs", 1)
    add_heading(doc, "12.1 Users and role-based access", 2)
    add_bullets(doc, [
        "Super Admin opens Users to create accounts, assign roles, edit position/designation, and configure menu or Add/Edit/Delete permissions.",
        "Role changes apply to accounts that use that role. Verify the user signs out and signs back in before testing a changed menu.",
        "Position/designation is used in report signatories such as Dean, Head of Unit, Administrative Aide, or IT Officer.",
    ])
    add_heading(doc, "12.2 Activity Logs", 2)
    add_para(doc, "Activity Logs record important actions such as creating, editing, issuing, checking, deleting, restoring, and changing schedules. Apply at least one filter before loading a large log list. Use date range and user filters to narrow the result.")
    add_callout(doc, "Privacy.", "Maintenance photos, signatures, staff names, and account data may be personal information. Share reports only with authorized recipients and follow your organization’s privacy policy and the Data Privacy Act of 2012.", fill="FFF6E5", border="E5A52B")

    add_heading(doc, "13. Backup, restore, and deployment checks", 1)
    add_para(doc, "Backup and restore is a Super Admin function. A backup is a copy of database data; it is not a replacement for a tested disaster-recovery plan.")
    add_heading(doc, "13.1 Automatic backup schedule", 2)
    add_bullets(doc, [
        "Open Database/Backup and Restore.",
        "Choose the monthly or weekly schedule option when enabled, then set the day and time.",
        "Save the schedule and confirm the status shows Configured.",
        "On Windows/XAMPP, keep the Laravel scheduler running with php artisan schedule:work, or run php artisan schedule:run from Windows Task Scheduler.",
        "A schedule test command can run the backup job immediately for verification. Check the configured backup directory afterward.",
    ])
    add_screenshot(doc, "backup-schedule.png", "Figure 12. Automatic backup schedule setup.", 6.2)
    add_screenshot(doc, "scheduler-running.png", "Figure 13. Example of the Laravel scheduler running on Windows.", 5.7)
    add_heading(doc, "13.2 Safe restore procedure", 2)
    add_bullets(doc, [
        "Download a fresh backup before restoring anything.",
        "Restore only a file from a trusted source and verify its date and size.",
        "Use a clone/test database first. Do not restore directly over production during business hours.",
        "After restore, test login, dashboard, equipment, checklist, reports, photos, and backup pages.",
        "If a restored database is missing newer columns, run the project migrations or use the deployment database upgrade procedure before accepting new data.",
        "Keep copies outside the web root and, ideally, on a separate drive or protected storage.",
    ])
    add_callout(doc, "Deployment safety.", "Disable directory listing, protect .env and storage directories, use HTTPS, keep APP_DEBUG=false in production, and expose only the public folder through the web server.", fill="FFF6E5", border="E5A52B")
    add_screenshot(doc, "deployment-error.png", "Figure 14. A deployment error example; use the troubleshooting checklist below rather than exposing this screen to users.", 5.5)

    add_heading(doc, "14. Troubleshooting and quick reference", 1)
    add_table(doc, ["What you see", "What to try first"], [
        ["The page is empty", "Press Apply/Generate or Reset. Large reports intentionally wait for a filter."],
        ["A filter shows no equipment", "Clear text search, choose the parent location before office, and reapply. Check status/condition filters."],
        ["Buttons do not react", "Reload once, confirm you are allowed to use the action, and check that no modal is open behind the page."],
        ["Camera opens a file picker", "Use HTTPS, allow camera access, and try a device with a rear camera. Upload is the fallback."],
        ["PDF/Excel is blank", "Verify filters, wait for the report to load, then preview before exporting."],
        ["A 500/server error appears", "Record the page and action, contact the system administrator, and do not expose the detailed error screen publicly."],
        ["Dark mode flashes white", "Clear cached assets and reload; use the official deployed URL with the correct built assets."],
        ["Backup scheduler is not running", "Check the scheduler process or Windows Task Scheduler and confirm the backup directory timestamp."],
    ], [2.2, 4.05])
    add_heading(doc, "Quick safety checklist", 2)
    add_bullets(doc, [
        "Use the correct location and office before saving equipment or schedules.",
        "Do not reuse a property number or delete history casually.",
        "Preview reports before printing or distributing them.",
        "Keep signatures, photos, and staff information confidential.",
        "Verify backups by restoring a copy in a test environment.",
        "Contact the PMAMS administrator or Contributors page for support.",
    ])
    add_para(doc, "End of manual • PMAMS Version 3.5.7 © 2026", italic=True, color=(89, 101, 114))

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
