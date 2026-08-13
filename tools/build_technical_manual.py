from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "manual_assets"
OUT = ROOT / "docs" / "PMAMS_Technical_User_Manual_v3.5.2.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "5B6770"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF8E5"
PALE_RED = "FCE8E6"
PALE_GREEN = "EAF5ED"
WHITE = "FFFFFF"
GRID = "AAB7C4"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para(p, before=0, after=6, line=1.25, align=None, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True


def add_text(doc, text, size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6):
    p = doc.add_paragraph()
    set_para(p, before=before, after=after, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(r, size=13, color=BLUE, bold=True)
    else:
        set_run_font(r, size=12, color=DARK_BLUE, bold=True)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    set_para(p, after=4)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    set_para(p, after=4)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    r = p.add_run(text)
    set_run_font(r, name="Consolas", size=9.5, color="28323C")
    return p


def add_note(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para(p, after=0, line=1.15)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, color=accent, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, header_fill=PALE_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent=120)
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        set_para(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(header)
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            set_para(p, after=0, line=1.08)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    return table


def add_figure(doc, path, caption, width=6.1):
    if not path.exists():
        add_note(doc, "Screenshot unavailable.", f"Expected asset {path.name} was not found.", fill=PALE_RED, accent="9B1C1C")
        return
    p = doc.add_paragraph()
    set_para(p, before=5, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    set_para(cap, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def set_header_footer(section):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    set_para(hp, after=0, line=1.0)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("PMAMS  |  Technical User Manual")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    set_para(fp, after=0, line=1.0)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("PMAMS v3.5.2  |  Internal operations reference  |  Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_number(fp)


def configure_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True
    # Normal and heading styles: compact reference guide preset.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True
    set_header_footer(sec)


def add_cover(doc):
    p = doc.add_paragraph()
    set_para(p, before=12, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    for filename in ("catsu-logo.png", "ictu-logo.png"):
        r = p.add_run()
        r.add_picture(str(ASSETS / filename), width=Inches(0.95))
        r.add_text("   ")
    add_text(doc, "PMAMS", size=30, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_text(doc, "Preventive Maintenance and Asset Monitoring System", size=16, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_text(doc, "Comprehensive Technical User Manual", size=22, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=4)
    add_text(doc, "Operations, administration, deployment, security, and troubleshooting reference", size=12, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_note(doc, "Document control.", "Version 3.5.2 | Prepared 13 August 2026 | Audience: Super Admins, administrators, unit heads/deans, custodians, support staff, and deployment operators.", fill=PALE_BLUE, accent=BLUE)
    add_text(doc, "This manual describes the deployed PMAMS workflows and technical controls. Menu visibility and write actions are role-controlled; do not use this document as a substitute for your institution's security, backup, or records-retention policy.", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=8)
    add_text(doc, "Confidential operational reference", size=9.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=40, after=0)
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "Contents", 1)
    entries = [
        "1. Purpose and system boundaries",
        "2. Architecture and data flow",
        "3. Prerequisites and deployment profiles",
        "4. Configuration and environment variables",
        "5. Installation, upgrade, and cache procedures",
        "6. Authentication, roles, and permissions",
        "7. Navigation, SPA behavior, and responsive UI",
        "8. Locations, offices, staff, and designations",
        "9. Equipment inventory and data model",
        "10. Issuance, return, reissue, relocation, and history",
        "11. Preventive maintenance checklist",
        "12. PM Plan scheduling and completion monitoring",
        "13. Reports, filters, PDF, and Excel exports",
        "14. Maintenance Attention and local recommendations",
        "15. Maintenance Photo Gallery",
        "16. Activity Logs and auditability",
        "17. Backup, restore, and scheduler operations",
        "18. Recycle Bin and permanent deletion",
        "19. Android wrapper deployment",
        "20. Security, performance, and QA checklist",
        "21. Troubleshooting and support runbook",
        "Appendix A. Common commands",
        "Appendix B. Data dictionary quick reference",
    ]
    for entry in entries:
        add_bullet(doc, entry)
    add_note(doc, "Reading convention.", "Steps marked Super Admin require the highest privilege. Examples use Windows/Laragon paths; replace them with the XAMPP or Linux path used by your deployment.", fill=PALE_GOLD, accent="7A5A00")


def build_body(doc):
    add_heading(doc, "1. Purpose and system boundaries", 1)
    add_text(doc, "PMAMS is a Laravel application for registering ICT equipment, organizing locations and offices, issuing assets to end users, recording preventive-maintenance checklists, producing audit-ready reports, and surfacing maintenance-attention recommendations. This manual covers both day-to-day administration and the technical procedures needed to deploy and operate the system safely.")
    add_text(doc, "The application is authoritative for the records stored in its database. Uploaded photos, generated reports, SQL backups, and the Android wrapper are separate artifacts and must be protected by the same backup and access-control policy.")
    add_note(doc, "Boundary.", "PMAMS does not replace the institution's accounting, procurement, HR, identity-management, or records-retention systems. Treat imported staff and location data as references and reconcile them with the source system before bulk import.", fill=PALE_GOLD, accent="7A5A00")

    add_heading(doc, "2. Architecture and data flow", 1)
    add_text(doc, "A typical deployment has a browser or Android WebView client, Apache/Laragon/XAMPP serving Laravel's public directory, PHP 8.2+, MySQL/MariaDB, and a scheduler process. Livewire handles interactive page sections while Laravel controllers, policies, models, and queued/report services apply authorization and persistence rules.")
    add_table(doc, ["Layer", "Responsibility", "Operational check"], [
        ("Browser / Android", "Authentication, responsive UI, camera/QR permissions, file uploads", "HTTPS, camera permission, stable base URL"),
        ("Laravel", "Routes, validation, roles, audit events, reports, imports", "APP_KEY set; caches cleared after deploy"),
        ("Database", "Equipment, assignments, checklists, PM Plans, users, settings, logs", "Migrations applied; backup verified"),
        ("Storage", "Equipment photos, maintenance gallery photos, signatures, generated files", "public storage link and write permissions"),
        ("Scheduler", "Monthly/weekly database backups and optional attention snapshots", "schedule:work or Task Scheduler running"),
    ], [1800, 4700, 2860])
    add_figure(doc, ASSETS / "contact-sheet.png", "Figure 1. PMAMS workflow surfaces captured from the deployed application.", width=6.2)

    add_heading(doc, "3. Prerequisites and deployment profiles", 1)
    add_heading(doc, "3.1 Supported local profiles", 2)
    add_bullet(doc, "Laragon: place the project under C:\\laragon\\www\\pms_systemv2, start Apache and MySQL, and use the Laragon virtual host or artisan serve.")
    add_bullet(doc, "XAMPP/Apache: point the virtual host DocumentRoot to the project's public directory. Do not publish the repository root; it exposes artisan, vendor, config, and source files.")
    add_bullet(doc, "Linux/Nginx or managed hosting: use PHP-FPM, point the server root to public, enable rewrite rules, and keep storage/app and .env outside web access.")
    add_heading(doc, "3.2 Required software and PHP extensions", 2)
    add_table(doc, ["Component", "Requirement", "Why it matters"], [
        ("PHP", "8.2 or newer", "Laravel 12 runtime"),
        ("Database", "MySQL 8+ or MariaDB", "Transactions, JSON fields, foreign keys"),
        ("PHP extensions", "intl, gd, zip, mbstring, openssl, pdo_mysql, fileinfo, tokenizer, xml, ctype", "Validation, images, Excel/ZIP export, PDF and encryption"),
        ("Composer / Node", "Composer 2; Node/npm for Vite builds", "Dependency installation and frontend compilation"),
        ("HTTPS", "TLS certificate in production", "Secure sessions, camera APIs, login and Android fallback"),
    ], [1800, 3200, 4360])
    add_note(doc, "Important.", "Browser camera access normally requires HTTPS (or localhost during development). A production app served only over HTTP may show a working page but still fail camera or secure-cookie behavior.", fill=PALE_RED, accent="9B1C1C")

    add_heading(doc, "4. Configuration and environment variables", 1)
    add_text(doc, "Copy .env.example to .env and set deployment-specific values. Never commit .env, APP_KEY, database passwords, reCAPTCHA secrets, or service-account credentials to GitHub.")
    add_table(doc, ["Setting", "Typical value", "Technical effect"], [
        ("APP_ENV", "production", "Controls environment-specific behavior"),
        ("APP_DEBUG", "false", "Prevents stack traces and secrets from reaching users"),
        ("APP_URL", "https://host.example/pmams/public", "URL generation, redirects, email and Android links"),
        ("DB_*", "MySQL/MariaDB connection", "Primary application database"),
        ("SESSION_DRIVER", "database or file", "Where authenticated sessions are stored"),
        ("SESSION_LIFETIME", "480", "Minutes before idle session expiration"),
        ("SESSION_SECURE_COOKIE", "true on HTTPS", "Sends cookies only over TLS"),
        ("SESSION_SAME_SITE", "lax", "Cross-site request protection while preserving normal navigation"),
        ("FILESYSTEM_DISK", "public", "Equipment/gallery photo storage"),
        ("RECAPTCHA_*", "production site/secret keys", "Optional layered login protection"),
    ], [2200, 2600, 4560])
    add_code(doc, "APP_ENV=production\nAPP_DEBUG=false\nAPP_URL=https://example.edu/pmams/public\nSESSION_DRIVER=database\nSESSION_LIFETIME=480\nSESSION_SECURE_COOKIE=true\nSESSION_SAME_SITE=lax\nFILESYSTEM_DISK=public")
    add_note(doc, "Cookie behavior.", "SESSION_SECURE_COOKIE=true is appropriate only when the public URL is HTTPS. On a plain HTTP LAN URL, a secure cookie will not be sent and login may appear to reset. Use HTTPS for deployment and keep SameSite=Lax unless an approved integration requires another policy.", fill=PALE_GOLD, accent="7A5A00")

    add_heading(doc, "5. Installation, upgrade, and cache procedures", 1)
    add_heading(doc, "5.1 First installation", 2)
    for step in [
        "Clone or copy the project into the server's application directory.",
        "Run composer install --no-dev --optimize-autoloader for production and npm ci followed by npm run build when frontend assets are part of the release.",
        "Create .env, generate APP_KEY, and configure DB_* values.",
        "Run php artisan migrate --seed (or the approved migration procedure for an existing database).",
        "Run php artisan storage:link and confirm public/storage resolves to storage/app/public.",
        "Run php artisan optimize:clear, then cache configuration/routes/views only after validating the environment.",
        "Set web-server DocumentRoot to public and verify /login, /storage, a report, and an image URL before importing live data.",
    ]:
        add_number(doc, step)
    add_heading(doc, "5.2 Upgrade procedure", 2)
    for step in [
        "Take a fresh SQL backup and verify it is non-empty before changing code or schema.",
        "Put the application in a maintenance window; stop heavy imports and scheduler jobs.",
        "Deploy code and dependencies, then run migrations and clear/rebuild caches.",
        "Smoke-test login, equipment filters, add/edit, checklist, PM Plan, reports, photo gallery, backup page, and role restrictions.",
        "Re-enable scheduler and monitor logs for the first backup and report generation.",
    ]:
        add_number(doc, step)
    add_code(doc, "php artisan down --render=errors::503\ncomposer install --no-dev --optimize-autoloader\nphp artisan migrate --force\nphp artisan storage:link\nphp artisan optimize:clear\nphp artisan optimize\nphp artisan up")
    add_note(doc, "Do not skip migrations.", "Restoring an older SQL file can remove newer columns such as deleted_at or signature fields. Always run the current migrations after importing a legacy database, and test the application before returning it to users.", fill=PALE_RED, accent="9B1C1C")

    add_heading(doc, "6. Authentication, roles, and permissions", 1)
    add_text(doc, "PMAMS uses authenticated user accounts with role-based menu and action permissions. The Super Admin can manage role defaults and per-user access attributes. A permission change should be tested in a fresh session or after clearing the current session so stale navigation does not hide a newly granted action.")
    add_table(doc, ["Role", "Typical authority", "Important boundaries"], [
        ("Super Admin", "All modules, users, role settings, backups/restores, recycle bin, destructive cleanup, PM Plan setup", "Protect account and recovery path; use only for administration"),
        ("Admin", "Equipment, locations, offices, staff, reports, assigned checklist work", "User administration and destructive actions depend on configured attributes"),
        ("Unit Head / Dean", "Review/sign-off, assigned reports and location/college context", "Does not automatically gain Super Admin operations"),
        ("Custodian", "Permitted location/office/staff/equipment tasks and reports", "Mark-as-checked and other actions can be disabled through role settings"),
    ], [1700, 4700, 2960])
    add_heading(doc, "6.1 Permission verification checklist", 2)
    for item in [
        "Create a test account for each role; never test destructive permissions with the production Super Admin account.",
        "Verify sidebar visibility and direct-URL protection. Hiding a menu entry is not authorization by itself.",
        "Test Add, Edit, Delete, Restore, Export, Mark Checked, PM Plan, Database, Users, Recycle Bin, and Checklist Cleanup separately.",
        "After changing role defaults, sign out/in and verify a newly loaded page; then test the same route in a second browser profile.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Navigation, SPA behavior, and responsive UI", 1)
    add_text(doc, "The application uses Livewire/SPA-style updates for many menu and filter actions. A section update should preserve the sidebar, current filters, focus, and camera state where possible. If a page appears stale, use Reset or a hard refresh only after checking the browser console and Laravel log.")
    add_figure(doc, ASSETS / "mobile-menu.png", "Figure 2. Responsive mobile navigation and role-filtered menu entries.", width=3.3)
    add_heading(doc, "7.1 Browser test matrix", 2)
    add_table(doc, ["Test", "Expected result", "Evidence"], [
        ("Desktop light/dark mode", "No white flash; sidebar and header retain theme during filter/navigation", "Screen recording or browser screenshot"),
        ("Mobile sidebar", "Hamburger opens/closes; selecting a route closes or preserves sidebar consistently", "Mobile viewport test"),
        ("Auto-search fields", "Typed value remains after query; focus is restored when query completes", "Equipment/issuance/attention filters"),
        ("Camera", "Permission prompt once; preview/capture/stop works without file picker substitution", "Checklist/gallery/equipment photo"),
        ("SPA navigation", "No duplicate notifications or stale scripts after repeated navigation", "Browser console and UI smoke test"),
    ], [2200, 5000, 2160])

    add_heading(doc, "8. Locations, offices, staff, and designations", 1)
    add_text(doc, "Location is the parent record. Office is a child of a location. Staff records belong to an office and may carry a Position/Designation. The designated Head of Unit or Dean is used in PM Schedule Monitoring and Maintenance Attention report signatories. Keep one active designation per office and update it when responsibility changes.")
    add_figure(doc, ASSETS / "offices.png", "Figure 3. Office records under a registered location.", width=6.2)
    add_figure(doc, ASSETS / "staff.png", "Figure 4. Staff records and position/designation context.", width=6.2)
    add_heading(doc, "8.1 Recommended data-entry order", 2)
    for step in [
        "Create the location/college.",
        "Create offices beneath that location.",
        "Add staff to each office and record Position/Designation.",
        "Set the office Head of Unit or college Dean reference.",
        "Only then import equipment or issue equipment so office and staff matching is deterministic.",
    ]:
        add_number(doc, step)

    add_heading(doc, "9. Equipment inventory and data model", 1)
    add_text(doc, "Equipment is identified by a unique property number. A parent system unit can own child peripherals through Part of Property Number. Child records keep their own unique property numbers but are grouped for checklist, export, and history purposes.")
    add_table(doc, ["Field group", "Examples / rule", "Used by"], [
        ("Identity", "Equipment Type, Property Number, Part of Property Number, Serial Number", "Inventory, QR, exports, checklist"),
        ("Specifications", "Computer Name, Brand, Model, Memory, Storage, Form Factor, OS/MS Office", "Inventory and Maintenance Attention"),
        ("Placement", "Location deployed, Office deployed, issued end user", "Location filters, issuance, reports"),
        ("Lifecycle", "Date Acquired, Condition, Status, Last Maintenance Date", "Dashboard, PM progress, attention scoring"),
        ("Evidence", "Maintenance remarks, equipment photo, activity history", "Audit trail and reports"),
    ], [1800, 4300, 3260])
    add_heading(doc, "9.1 Supported equipment-specific behavior", 2)
    add_bullet(doc, "Desktop and Laptop: eligible for the full maintenance checklist and hardware/OS attention rules.")
    add_bullet(doc, "Monitor, UPS, AVR, Printer, Scanner, and Network Device: may be linked as peripherals; checklist eligibility and status rules depend on the current workflow configuration.")
    add_bullet(doc, "Network Device: select Access point, Router, Switch (managed), or Switch (unmanaged); record MAC Address and link Location Deployed/Office.")
    add_bullet(doc, "Memory: use the controlled values 2GB, 4GB, 8GB, 16GB, 32GB, or 64GB where available.")
    add_bullet(doc, "Storage: choose SSD or HDD, then select a capacity. Do not type a free-form capacity if a controlled option is available.")
    add_figure(doc, ASSETS / "equipment-filters.png", "Figure 5. Inventory filter row and search behavior.", width=6.2)
    add_heading(doc, "9.2 Import controls", 2)
    for item in [
        "Use the current XLSX template and keep the header names unchanged.",
        "Blank or zero property numbers follow the current temporary-number policy; rows without the required parent/child linkage may be skipped or rejected according to the import validator.",
        "Staff and office values must match active records. Import does not silently create trustworthy staff identities from free text.",
        "Run a small preview import first, review row errors, then import the validated file. Preserve the original source file as evidence.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. Issuance, return, reissue, relocation, and history", 1)
    add_text(doc, "Issuance assigns an available equipment record to an active end user. Reissue changes the assigned end user and updates the user's registered office. Relocation changes a location/office reference with an activity remark. The system keeps assignment history so a report can show origin and destination context.")
    add_figure(doc, ASSETS / "staff-equipment.png", "Figure 6. Equipment currently assigned to a staff member.", width=6.2)
    add_heading(doc, "10.1 Safe transaction sequence", 2)
    for step in [
        "Search by property number, serial number, model, or staff/office.",
        "Confirm the equipment is not already assigned or condemned.",
        "Select the end user; use the auto-search result rather than free text.",
        "Enter an activity remark when the action changes ownership, location, or operational intent.",
        "Save and verify the equipment detail page, staff equipment list, dashboard count, and activity log.",
    ]:
        add_number(doc, step)

    add_heading(doc, "11. Preventive maintenance checklist", 1)
    add_text(doc, "The checklist records the state of a system unit and its linked peripherals for a specific maintenance date. Each checklist row must be resolved according to the equipment type: OK, Not OK, or Not Available where supported. A Not OK disposition can expose Repair or Condemn/Not in Use controls, depending on the configured checklist version.")
    add_heading(doc, "11.1 Checklist workflow", 2)
    for step in [
        "Open Mark as Checked for an eligible Desktop or Laptop equipment record.",
        "Confirm the parent property number and linked peripherals. Use the link shortcut for a missing child; do not create a duplicate parent.",
        "Select one outcome per checklist item. Resolve Not OK dispositions and provide remarks/corrective action when required.",
        "Use Not Available only for supported peripheral checks. The system may generate default remarks/corrective action for UPS/AVR or printer rules.",
        "Capture a maintenance photo from the device camera or upload a valid image; verify it appears in PM Gallery after saving.",
        "Review the generated remarks, condition, status, and current linked-property values before saving.",
        "If duplicate detection is triggered within the configured window, verify the existing checklist and enter the requested reason only when the workflow asks for it.",
    ]:
        add_number(doc, step)
    add_note(doc, "Historical records.", "Checklist entries are historical events. Deleting or restoring a checklist through the approved cleanup workflow should not overwrite another maintenance date. Keep the checklist history and equipment Last Maintenance Date synchronized according to the current retention policy.", fill=PALE_GOLD, accent="7A5A00")

    add_heading(doc, "12. PM Plan scheduling and completion monitoring", 1)
    add_text(doc, "PM Plan defines the maintenance schedule by location and office. Super Admin publishes original schedules and assigns one or more Admin/Super Admin accounts. Admins complete the checklist for equipment in assigned locations. An override schedule is stored alongside the original schedule with a reason; it does not erase the original plan.")
    add_figure(doc, ASSETS / "pm-plan-card.png", "Figure 7. PM Plan card showing schedule, assignment, progress, and actions.", width=6.2)
    add_figure(doc, ASSETS / "pm-override.png", "Figure 8. PM Plan override form and reason tracking.", width=6.2)
    add_heading(doc, "12.1 Duplicate detection and completion", 2)
    add_bullet(doc, "Duplicate detection uses location, office, schedule period, and existing plan records; recycled plans are also considered by the current implementation.")
    add_bullet(doc, "The completion percentage is based on eligible, non-condemned equipment records and saved checklist history. If equipment is transferred or condemned, validate that the target scope and counts are the expected ones.")
    add_bullet(doc, "Completion details include actual date, signer/recorded names, signature pad data where used, and remarks. These details feed PM Schedule Monitoring reports.")

    add_heading(doc, "13. Reports, filters, PDF, and Excel exports", 1)
    add_text(doc, "Reports are intentionally filter-first for large data sets. Apply location and office filters together when office is a child of location. When a report says it remains unloaded until Apply or Reset, that is an intentional performance guard.")
    add_table(doc, ["Report", "Primary use", "Verification points"], [
        ("Checked Equipment", "Maintenance history by date, equipment, location, remarks, corrective action", "Selected rows, PDF preview, cleanup permissions"),
        ("PM Schedule Monitoring", "Original/override schedules, actual maintenance dates, sign-off", "Office mapping, override reason, signature placement"),
        ("Maintenance Attention", "Prioritized equipment needing upgrade/procurement/maintenance", "Location/office/type/status filters, score and priority"),
        ("Quality Objective Monitoring", "Baseline vs actual maintenance performance by semester", "PM Plan progress, transfer adjustments, MET/UNMET formula"),
        ("All Assets / Issuance", "Inventory and assignment exports", "Filtered data, issued and unissued records, origin/destination"),
    ], [2200, 4300, 2860])
    add_figure(doc, ASSETS / "maintenance-attention-report.png", "Figure 9. Maintenance Attention report output with header, data table, and footer.", width=6.2)
    add_heading(doc, "13.1 Export validation", 2)
    for item in [
        "Confirm the filter summary at the top of the report matches the controls used.",
        "Open the PDF and the Excel workbook from the same filter state; the records and calculations must agree.",
        "Check long text columns for wrapping, no clipped rows, repeated headers on continuation pages, and footer/header placement.",
        "Verify blank values remain blank where the report contract requires it; do not infer a hyphen as a real value.",
        "For Excel, review formulas and cached values in the performance/quality objective sheet before distribution.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "14. Maintenance Attention and local recommendations", 1)
    add_text(doc, "Maintenance Attention is a rule-based/offline recommendation surface. The current logic covers configured Desktop/Laptop hardware and can include Printer, Monitor, and UPS condition/status checks. It can flag low RAM, HDD storage, unsupported/old operating systems, cracked licenses, and age thresholds. Priority and score are explanations for review, not an automatic purchase order.")
    add_table(doc, ["Signal", "Example recommendation", "Operator action"], [
        ("Age", "Equipment is at least six years old", "Validate lifecycle, repair history, and procurement plan"),
        ("Memory", "Windows 10/11 with RAM at or below 8GB", "Consider upgrade to at least 16GB"),
        ("Legacy OS", "Windows 8 or below with RAM at or below 4GB", "Consider upgrade to at least 8GB and supported OS"),
        ("Storage", "Desktop uses HDD", "Consider SSD flash storage upgrade"),
        ("License", "Cracked or non-genuine OS/MS Office", "Procure genuine OS and Microsoft suite"),
        ("Condition/status", "Unserviceable with Available/Repair/Not in Use status", "Confirm condition, disposition, and history"),
    ], [1800, 4300, 3260])
    add_note(doc, "Model mode.", "The system can expose Laravel-coded rules and an optional offline local model mode. If a model is unavailable, the rule engine should remain the safe fallback. Always show the criteria, model date, confidence, and source mode in the UI/report so a reviewer can challenge a recommendation.", fill=PALE_GREEN, accent="1D6B42")
    add_heading(doc, "14.1 Model operations", 2)
    add_bullet(doc, "Record the model training date/time and data window when a local model is retrained.")
    add_bullet(doc, "Use historical checklist/equipment records only after confirming the retention, privacy, and quality policy.")
    add_bullet(doc, "Treat ONNX/scikit-learn output as a ranking signal. The final maintenance decision remains with the responsible office/ICT unit.")

    add_heading(doc, "15. Maintenance Photo Gallery", 1)
    add_text(doc, "PM Gallery stores preventive-maintenance photos with optional equipment and checklist links, capture date/time, caption, and uploader ownership. Users may delete their own uploaded photos; Super Admin policy may permit broader cleanup. Bulk download packages selected images into a ZIP file.")
    add_figure(doc, ASSETS / "maintenance-table.png", "Figure 10. Maintenance photo/checklist evidence surface.", width=6.2)
    add_bullet(doc, "Maximum file size is 10 MB per photo; accept only supported image types.")
    add_bullet(doc, "Prefer the rear-camera Take Photo action on mobile. If the browser shows a file picker instead, verify HTTPS, camera permission, and input capture attributes.")
    add_bullet(doc, "After saving from the checklist, confirm the photo appears in PM Gallery and is linked to the checklist/equipment record.")

    add_heading(doc, "16. Activity Logs and auditability", 1)
    add_text(doc, "Activity Logs record important create, update, delete, restore, export, assignment, checklist, PM Plan, and backup actions. The log is not a replacement for database backups; it is an operational audit trail. Use action, subject type, user, and date-range filters, and avoid loading the complete log in a high-volume production database before a filter is applied.")
    add_table(doc, ["Audit question", "Where to look"], [
        ("Who changed equipment or assignment?", "Activity Logs; equipment history; issuance history"),
        ("Why was a PM schedule changed?", "PM Plan override reason and Activity Logs"),
        ("Who restored or permanently deleted a record?", "Recycle Bin and Activity Logs"),
        ("Was a report/export generated?", "Activity Logs and output filename/time"),
    ], [3600, 5760])

    add_heading(doc, "17. Backup, restore, and scheduler operations", 1)
    add_text(doc, "Database backup is a Super Admin operation. The backup page supports manual SQL download, a configurable monthly/weekly schedule, and restore with a pre-restore safety copy. A restore may replace application data and is not fully transactional at the database level; test in a clone before production use.")
    add_figure(doc, ASSETS / "backup-schedule.png", "Figure 11. Automatic backup schedule configuration.", width=6.2)
    add_figure(doc, ASSETS / "scheduler-running.png", "Figure 12. Windows process showing artisan schedule:work.", width=6.2)
    add_heading(doc, "17.1 Manual backup test", 2)
    for step in [
        "Open Database / Backup and Restore as Super Admin.",
        "Download a fresh SQL backup and record its timestamp and size.",
        "Inspect the SQL header and confirm it contains all expected base tables, including newer tables/columns after migrations.",
        "Restore the file only in a cloned database first; verify login, equipment, PM Plan, checklist, photos, users, and settings.",
        "Keep the pre-restore safety file until the clone passes validation.",
    ]:
        add_number(doc, step)
    add_heading(doc, "17.2 Scheduler", 2)
    add_code(doc, "php artisan schedule:list\nphp artisan schedule:test --name=database:backup-monthly --no-interaction\nphp artisan schedule:work")
    add_bullet(doc, "schedule:work is a long-running process. On Windows, register it in Task Scheduler or run it as a service under a controlled account.")
    add_bullet(doc, "Check the process command line, Laravel logs, backup directory, and last modified timestamp after the scheduled time.")
    add_bullet(doc, "The scheduler should not overwrite existing backup files; use timestamped names and a retention policy.")
    add_note(doc, "Restore warning.", "Never import an unverified SQL file over a live database. A syntax/packet/foreign-key failure can leave a partial restore. Use the clone-restore test and current migrations before deployment.", fill=PALE_RED, accent="9B1C1C")

    add_heading(doc, "18. Recycle Bin and permanent deletion", 1)
    add_text(doc, "Soft deletion protects users, equipment, locations, and PM Plans from accidental removal. Restore selected records when possible. Permanent delete is irreversible and should be reserved for Super Admin after confirming the record, dependencies, export evidence, and retention requirements.")
    add_table(doc, ["Action", "Expected behavior", "Control"], [
        ("Move to recycle bin", "Record is hidden from normal active lists but retained", "Role permission and confirmation"),
        ("Restore", "Record returns to active data with links where valid", "Check duplicate/foreign-key conflicts"),
        ("Permanent delete", "Record and related files/history may be erased", "Super Admin only; confirm scope"),
        ("Bulk delete", "Selected or filtered pages are processed", "Select-all scope, filters, and confirmation"),
    ], [2400, 4560, 2400])

    add_heading(doc, "19. Android wrapper deployment", 1)
    add_text(doc, "The Android app is a WebView wrapper around PMAMS. Keep the local/LAN URL and public HTTPS fallback configured for the deployment environment. The app should ask for confirmation before switching network endpoints, then suppress repeated prompts while the active endpoint remains reachable.")
    add_bullet(doc, "Build from the android-app project using the bundled Android/Gradle toolchain configured for the release.")
    add_bullet(doc, "Set versionCode/versionName deliberately for each release; versioning is not automatically meaningful unless the build files are updated.")
    add_bullet(doc, "Test login, camera, QR scanner, file upload, SPA navigation, fallback URL, and back-button behavior on at least one older and one current Android device.")
    add_bullet(doc, "Distribute the APK only through an approved channel and keep the signing key out of the repository.")

    add_heading(doc, "20. Security, performance, and QA checklist", 1)
    add_heading(doc, "20.1 Security baseline", 2)
    for item in [
        "APP_DEBUG=false in production; errors go to logs, not browser stack traces.",
        "HTTPS enforced; secure session cookies enabled; HSTS and secure headers configured at the web server.",
        "Repository root blocked from directory listing; only public is web-accessible.",
        "Uploads limited by size/type and served through controlled storage routes where needed.",
        "Authorization tested server-side for direct URLs, exports, delete, restore, backup, and role management.",
        "Secrets, APP_KEY, backup files, logs, and user data excluded from Git and public storage.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "20.2 Performance/load baseline", 2)
    add_bullet(doc, "Use filter-first loading for equipment, activity logs, reports, maintenance attention, and gallery lists.")
    add_bullet(doc, "Use pagination and bounded exports; do not load every row into a browser or PHP array for ordinary page views.")
    add_bullet(doc, "Cache stable lookups such as equipment types and locations, but invalidate caches after changes.")
    add_bullet(doc, "Monitor slow queries, PHP memory, max execution time, upload limits, and database packet size for imports/restore.")
    add_bullet(doc, "Test with realistic equipment, checklist, photo, and activity-log volumes before launch.")
    add_heading(doc, "20.3 Release smoke test", 2)
    add_table(doc, ["Area", "Pass criteria"], [
        ("Login/session", "Login, remember-me, timeout, second-tab redirect, logout"),
        ("Equipment", "Add/edit/view, generated number, child link, filters, photo, QR, export"),
        ("Locations", "Location/office/staff add/edit/delete, head/dean designation, child office filter"),
        ("Maintenance", "Checklist validation, duplicate window, photo/gallery, history, PM progress"),
        ("Reports", "Filtered PDF/Excel values, headers/footers, long text wrapping, pagination"),
        ("Recovery", "Backup download, clone restore, recycle restore, permanent-delete permissions"),
    ], [2200, 7160])

    add_heading(doc, "21. Troubleshooting and support runbook", 1)
    add_table(doc, ["Symptom", "Likely cause", "First response"], [
        ("500 error after deployment", "Missing migration, extension, APP_KEY, permissions, or cached config", "Check storage/logs/laravel.log, run optimize:clear, verify migrations/extensions"),
        ("Login resets on HTTPS", "SESSION_SECURE_COOKIE/SAME_SITE/APP_URL mismatch", "Set HTTPS APP_URL, secure cookie true, clear config/session cache"),
        ("Camera opens file picker", "Insecure origin or camera permission/capture attribute issue", "Use HTTPS/localhost, grant permission, test supported browser"),
        ("Photo broken after XAMPP transfer", "storage link/path or missing file", "Run storage:link, verify storage/app/public and route URL"),
        ("Report blank or clipped", "Unapplied filters, missing variables, PDF CSS/table widths", "Check response/logs, inspect filter summary, render PDF, verify view variables"),
        ("Restore fails part-way", "SQL packet, foreign key, syntax, or schema mismatch", "Stop; keep pre-restore safety file; restore clone and run current migrations"),
        ("Scheduler did not run", "schedule:work stopped or Task Scheduler misconfigured", "Check process command line, schedule:list, logs, backup timestamp"),
    ], [2500, 3500, 3360])
    add_heading(doc, "21.1 Evidence to collect before escalation", 2)
    add_bullet(doc, "URL, role, date/time, browser/device, and exact action sequence.")
    add_bullet(doc, "Screenshot or sanitized error message, Laravel log excerpt, and activity-log entry ID if available.")
    add_bullet(doc, "Relevant filter values, property number/office/location, and whether the issue occurs in another browser/profile.")
    add_bullet(doc, "For deployment failures: PHP version/extensions, database version, APP_URL, storage link state, and last successful migration.")

    add_heading(doc, "Appendix A. Common commands", 1)
    add_code(doc, "php artisan about\nphp artisan route:list\nphp artisan migrate:status\nphp artisan schedule:list\nphp artisan storage:link\nphp artisan optimize:clear\nphp artisan config:cache\nphp artisan route:cache\nphp artisan view:cache\nphp artisan test")
    add_text(doc, "Run commands from the project root. In production, use --force where Laravel requires explicit confirmation, and record the command/time in the deployment log.")

    add_heading(doc, "Appendix B. Data dictionary quick reference", 1)
    add_table(doc, ["Term", "Meaning"], [
        ("Property Number", "Unique inventory identifier for an equipment record"),
        ("Part of Property Number", "Parent property number used to group a peripheral with a system unit"),
        ("Condition", "Physical/operational condition such as Serviceable or Unserviceable"),
        ("Status", "Lifecycle/assignment state such as Available, Issued, Repair, or Not in Use"),
        ("Checklist history", "Immutable maintenance event record for a date and checker"),
        ("PM Plan", "Published preventive maintenance schedule for location/office scope"),
        ("Override", "Temporary alternate schedule stored with reason beside the original"),
        ("Maintenance Attention", "Rule/model-ranked list of equipment needing review"),
        ("Activity Log", "Audit record of a system action, actor, subject, and description"),
    ], [2600, 6760])
    add_note(doc, "Document maintenance.", "Update this manual whenever routes, roles, database schema, backup behavior, Android endpoints, or report contracts change. Pair each revision with a smoke-test record and a fresh screenshot set.", fill=PALE_BLUE, accent=BLUE)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    doc.add_page_break()
    build_body(doc)
    # Set core properties without personal user data.
    props = doc.core_properties
    props.title = "PMAMS Comprehensive Technical User Manual"
    props.subject = "Operations, administration, deployment, security, and troubleshooting reference"
    props.author = "PMAMS Documentation Team"
    props.keywords = "PMAMS, Laravel, preventive maintenance, equipment, deployment, technical manual"
    props.comments = "Generated from the PMAMS application documentation and checked-in screenshots."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
